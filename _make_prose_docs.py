"""Generate two DOCX files of the new v2-expansion prose from this session:
   1) AS-WRITTEN — exactly what is in the .twee (.claude-draft blocks).
   2) REWRITE — same content, stripped of recursive/paradoxical mannerisms.
   Each block is prefaced by 1–2 sentences saying what the block is for.
"""
from docx import Document
from docx.shared import Pt, RGBColor

# ---------------------------------------------------------------- pieces
PIECES = [
    {
        "title": "The Synthesis — opening",
        "purpose": (
            "Fires the moment the player clicks 'Perform the synthesis' at the "
            "Pillars of Hercules — the entry to the 6th-world ritual. Sets the "
            "scene: the three pillars, the five gifts in hand, the threshold."
        ),
        "as_written": (
            "You stand between the three pillars. The two broken ones face the "
            "third. The third is whole. The third is yours.\n\n"
            "You have everything. The mantra. The tracing. The rubbing. The "
            "number. The vision. Five gifts gathered across five nights. You "
            "did not know you were preparing.\n\n"
            "You begin."
        ),
        "rewrite": (
            "You stand between the three pillars. The two broken ones face the "
            "third, which is whole and yours.\n\n"
            "You have everything you need: the mantra, the tracing, the "
            "rubbing, the number, the vision. Five gifts gathered across five "
            "nights, without knowing you were preparing."
        ),
    },
    {
        "title": "Synthesis — gift 1 (mantra)",
        "purpose": (
            "Shows on click of 'Speak the mantra'. The MANTRA box appears above "
            "this prose, listing the twelve Vajra Guru syllables."
        ),
        "as_written": (
            "You say it, twelve syllables, into the air. The third pillar warms."
        ),
        "rewrite": (
            "You speak the twelve syllables into the air. The third pillar warms."
        ),
    },
    {
        "title": "Synthesis — gift 2 (tracing)",
        "purpose": (
            "Shows on click of 'Hold up the tracing'. The TRACING box shows "
            "'The Hummingbird, from above' (the Nazca gift)."
        ),
        "as_written": (
            "You hold the page out. The bird flies into the warm pillar and is "
            "held there."
        ),
        "rewrite": (
            "You hold the page out. The hummingbird crosses into the warm "
            "pillar and stays there."
        ),
    },
    {
        "title": "Synthesis — gift 3 (rubbing)",
        "purpose": (
            "Shows on click of 'Show the rubbing'. The RUBBING box shows "
            "'A glyph in Rongorongo' (the Easter Island gift)."
        ),
        "as_written": (
            "You show the rubbing. The glyph speaks itself, in a voice you can "
            "hear in your bones."
        ),
        "rewrite": (
            "You hold up the rubbing, and the glyph reads itself aloud. You "
            "feel the sound as much as hear it."
        ),
    },
    {
        "title": "Synthesis — gift 4 (number)",
        "purpose": (
            "Shows on click of 'Hum the number'. The PROPORTION box shows "
            "'A resonant number' (the Pyramid gift)."
        ),
        "as_written": (
            "You hum the number. The chamber that gave it to you comes with "
            "you, briefly. The pillars sit inside it."
        ),
        "rewrite": (
            "You hum the number, and the King's Chamber seems to surround you "
            "for a moment. The three pillars fit inside it."
        ),
    },
    {
        "title": "Synthesis — gift 5 (vision) + hexagram reveal",
        "purpose": (
            "Shows on click of 'Draw the wheel'. The VISION box shows 'The "
            "wheel-within-a-wheel'. This is the climax of the ritual — the "
            "pillars rearrange, the lily-pentangle joins them, and the "
            "Solomon's-Seal hexagram materialises (the gold SVG with the two "
            "interlocking pentangles and ten luminous outer points)."
        ),
        "as_written": (
            "You draw the wheel in the air. The eyes around its rim open. "
            "They watch you finish.\n\n"
            "Then the pillars rearrange. The two broken ones rise, whole. The "
            "third lifts. The three stand for an instant in a triangle, and "
            "from beneath them rises another — the lily-pentangle of Soho, "
            "which has been waiting for you in the map.\n\n"
            "The two figures meet."
        ),
        "rewrite": (
            "You draw the wheel in the air. The eyes around its rim open and "
            "watch you to the end of the gesture.\n\n"
            "The pillars rearrange themselves. The two broken ones stand "
            "whole; the third lifts. For an instant the three form a "
            "triangle, and beneath them rises the lily-pentangle of Soho "
            "that has been on the map all along.\n\n"
            "The two shapes lock into each other."
        ),
    },
    {
        "title": "Synthesis — closing on the Seal",
        "purpose": (
            "Shows immediately after the hexagram SVG materialises. Names what "
            "the player is looking at and opens the doorway through to The "
            "Sanctum (the 6th world)."
        ),
        "as_written": (
            "Solomon's Seal. Two pentangles, ten points, one figure. It is "
            "yours now.\n\n"
            "The figure opens like a door. You step into it."
        ),
        "rewrite": (
            "Solomon's Seal: two pentangles laid into one figure, ten points "
            "in all. Yours now.\n\n"
            "It opens like a door, and you step through."
        ),
    },
    {
        "title": "The Sanctum — first sight",
        "purpose": (
            "First passage on the far side of the hexagram. The slate-eyed "
            "guide who has been recurring across all five worlds is finally "
            "fully met. He has the player's book on the table. Includes the "
            "keystone time-stratigraphy beat (predecessors who carried the "
            "book before)."
        ),
        "as_written": (
            "You stand in a room that is not a room. The light is white but "
            "it is also gold. There is a long table. There is a chair. There "
            "is the man.\n\n"
            "He is not surprised to see you. He has been waiting for as long "
            "as you have been arriving.\n\n"
            "His eyes are slate, and kind.\n\n"
            "'You are not the first,' he says. 'You will not be the last. "
            "But you are the one who came tonight.'\n\n"
            "'Before you, a girl in Cairo with a satchel of charcoal. Before "
            "her, a sailor at Knossos. Before him, a priestess at a hilltop "
            "that is now buried. The book keeps coming back to me.'\n\n"
            "He gestures at the table. There is the book — your book. He has "
            "read it. He has been waiting for you to finish it.\n\n"
            "'It's ready,' he says."
        ),
        "rewrite": (
            "You're in a wide, still room. The light is a pale gold. A long "
            "table. A chair. And the man, at the chair.\n\n"
            "He looks up without surprise. He's been waiting a long time.\n\n"
            "His eyes are slate, and kind.\n\n"
            "'You're not the first to come,' he says. 'You won't be the "
            "last. But tonight it's you.'\n\n"
            "'A girl in Cairo brought me one version of this. A sailor at "
            "Knossos brought another. There was a priestess in a place that "
            "no longer exists. The book keeps coming back.'\n\n"
            "He nods toward the table. Your book is on it. He's read it, "
            "and he's been waiting for you to finish.\n\n"
            "'It's ready,' he says."
        ),
    },
    {
        "title": "The Sanctum — Sitting (middle)",
        "purpose": (
            "Player clicks 'Sit down'. Short transitional passage — they sit, "
            "the man sits opposite, the book opens to a page. Links onward "
            "to Alt-Dawn."
        ),
        "as_written": (
            "You sit. He sits opposite. The book between you.\n\n"
            "He opens it to a page you don't remember writing — but you "
            "wrote it.\n\n"
            "It is the dawn."
        ),
        "rewrite": (
            "You sit. He takes the chair opposite, the book between you.\n\n"
            "He opens it to a page you'd forgotten writing. You wrote it "
            "anyway.\n\n"
            "It opens onto the dawn."
        ),
    },
    {
        "title": "Alt-Dawn",
        "purpose": (
            "The alternate dawn ending — fires when the player has performed "
            "the synthesis (5/5 lifetime gifts). Currently a stub that flags "
            "where the alternate-recipient ending should live; Dr Quill to "
            "write the actual ending here. Includes an explicit author note "
            "to him in parentheses, which should be deleted on rewrite."
        ),
        "as_written": (
            "You read what you wrote.\n\n"
            "The book is finished. It has been finished for some time, "
            "though you didn't know.\n\n"
            "In v1's dawn the book went one way. In this dawn it goes "
            "another. (Dr Quill: this is where the alternate recipient "
            "lives — Aoife, an unnamed publisher, the slate-eyed man "
            "himself, the protagonist's own dying or rising form. Write "
            "the ending here that v1 cannot give.)\n\n"
            "The light is white now. It is also gold. The third pillar, in "
            "the room behind you, leans against itself.\n\n"
            "You finish reading. You close the book. You give it back."
        ),
        "rewrite": (
            "You read what you wrote.\n\n"
            "The book has been finished for some time. You hadn't known.\n\n"
            "[Stub note for you: in v1 the dawn delivered the book one way. "
            "This dawn should deliver it differently — to Aoife, an "
            "unnamed publisher, the slate-eyed man himself, the "
            "protagonist's dying or rising form. The ending v1 couldn't "
            "give. Write the body of the ending you want here.]\n\n"
            "The light has changed to a pale gold. The third pillar stands "
            "quietly behind you.\n\n"
            "You finish reading, close the book, and hand it back."
        ),
    },
    {
        "title": "The Cave (Himalayas) — stratigraphy beat",
        "purpose": (
            "Inserted into The Cave passage (the Himalayan world's centre, "
            "where the player meets the slate-eyed man and receives the "
            "second half of the mantra). One added sentence positions the "
            "transmission across deep time."
        ),
        "as_written": (
            "'I gave the first half to a Tibetan monk in 1124,' he says. 'He "
            "gave it to a Greek captain before that. The captain gave it to "
            "a Sumerian. Tonight it's yours.'"
        ),
        "rewrite": (
            "'A Tibetan monk had the first half of this before you,' he "
            "says. 'And a Greek captain before him. The chain goes back "
            "farther than I can remember. Tonight you have it.'"
        ),
    },
    {
        "title": "The Glyph (Easter Island) — stratigraphy beat",
        "purpose": (
            "Inserted into The Glyph (Easter Island centre) — the player has "
            "just taken a charcoal rubbing of the inside of a moai's mouth "
            "and is holding a glyph in Rongorongo. One added sentence pushes "
            "the script's age back beyond the moai themselves."
        ),
        "as_written": (
            "The moai were carved to remember this script. The script was "
            "older than them."
        ),
        "rewrite": (
            "The moai were carved to remember this script, which was "
            "already old when they were."
        ),
    },
]


# ---------------------------------------------------------------- helpers
def write_doc(path: str, version: str) -> None:
    doc = Document()

    # base style
    style = doc.styles["Normal"]
    style.font.name = "Garamond"
    style.font.size = Pt(12)

    title = doc.add_heading(
        f"Dream Street Shuffle — v2-expansion new prose ({version})",
        level=0,
    )

    intro = doc.add_paragraph()
    intro.add_run(
        "All the prose written during the v2-expansion build of 2026-05-14: "
        "the synthesis ritual, The Sanctum (6th world), Alt-Dawn (alternate "
        "ending stub), and the three deep-time stratigraphy beats inserted "
        "into The Cave, The Glyph, and The Sanctum. Each block is preceded "
        "by a short note saying what it is for."
    ).italic = True

    if version == "AS-WRITTEN":
        intro2 = doc.add_paragraph()
        r = intro2.add_run(
            "This file shows the prose exactly as it currently sits in the "
            ".twee, wrapped in the .claude-draft marker. The companion file "
            "(REWRITE) is a second pass that strips the recursive/"
            "paradoxical mannerisms you flagged."
        )
        r.italic = True
    else:
        intro2 = doc.add_paragraph()
        r = intro2.add_run(
            "This is the rewrite pass. Same blocks in the same order, but "
            "without the 'X is not X but X-er than X' construction, the "
            "stacked-short-clause rhythm, and the every-block closing-image "
            "moral. The point is to give you something flatter and less "
            "mannered to react against."
        )
        r.italic = True

    key = "as_written" if version == "AS-WRITTEN" else "rewrite"

    for piece in PIECES:
        doc.add_heading(piece["title"], level=1)
        # purpose note
        p = doc.add_paragraph()
        run = p.add_run(piece["purpose"])
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        # prose
        for para in piece[key].split("\n\n"):
            doc.add_paragraph(para)

    doc.save(path)


# ---------------------------------------------------------------- run
out_dir = "/Users/samquill/Claude work/Dream Street Shuffle - Game Files"
write_doc(f"{out_dir}/v2-prose-AS-WRITTEN.docx", "AS-WRITTEN")
write_doc(f"{out_dir}/v2-prose-REWRITE.docx", "REWRITE")
print("Wrote v2-prose-AS-WRITTEN.docx and v2-prose-REWRITE.docx")
